from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from database import get_db
from models import Document, User
from auth import SECRET_KEY, ALGORITHM
from jose import jwt, JWTError
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from services.summarizer import summarize
from services.ner import extract_entities
from services.rag import answer_question
from pydantic import BaseModel
import io

router = APIRouter()

# ── Get current user from token ──────────────────────
def get_current_user(token: str, db: Session):
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = int(payload.get("sub"))
        user = db.query(User).filter(User.id == user_id).first()
        if not user:
            raise HTTPException(status_code=401, detail="User not found")
        return user
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")

# ── Extract text from file ───────────────────────────
def extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.endswith(".pdf"):
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in pdf:
            text += page.get_text()
        return text

    elif filename.endswith(".docx"):
        docx = DocxDocument(io.BytesIO(file_bytes))
        return "\n".join([para.text for para in docx.paragraphs])

    elif filename.endswith(".txt"):
        return file_bytes.decode("utf-8")

    else:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, and TXT files supported")

# ── Routes ───────────────────────────────────────────
@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    token: str = None,
    db: Session = Depends(get_db)
):
    if not token:
        raise HTTPException(status_code=401, detail="Token required")

    user = get_current_user(token, db)
    file_bytes = await file.read()
    raw_text = extract_text(file_bytes, file.filename)

    doc = Document(
        user_id=user.id,
        filename=file.filename,
        raw_text=raw_text,
        summary=summarize(raw_text)
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {
        "message": "File uploaded successfully",
        "document_id": doc.id,
        "filename": doc.filename,
        "text_preview": raw_text[:300]
    }

@router.get("/documents")
def get_documents(token: str, db: Session = Depends(get_db)):
    user = get_current_user(token, db)
    docs = db.query(Document).filter(Document.user_id == user.id).all()
    return [
        {
            "id": d.id,
            "filename": d.filename,
            "summary": d.summary,
            "created_at": d.created_at
        }
        for d in docs
    ]
@router.get("/documents/{doc_id}/entities")
def get_entities(doc_id: int, token: str, db: Session = Depends(get_db)):
    user = get_current_user(token, db)
    
    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.user_id == user.id
    ).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    entities = extract_entities(doc.raw_text)
    return {
        "document_id": doc_id,
        "filename": doc.filename,
        "entities": entities
    }
class ChatRequest(BaseModel):
    question: str

@router.post("/documents/{doc_id}/chat")
def chat_with_document(
    doc_id: int,
    request: ChatRequest,
    token: str,
    db: Session = Depends(get_db)
):
    user = get_current_user(token, db)

    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.user_id == user.id
    ).first()

    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if not doc.raw_text:
        raise HTTPException(status_code=400, detail="Document has no text")

    answer = answer_question(request.question, doc.raw_text)
    return {
        "document_id": doc_id,
        "question": request.question,
        "answer": answer
    }